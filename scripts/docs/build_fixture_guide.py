"""Build the Word companion for docs/cms-fixture-dataset-guide.md."""

from pathlib import Path
import os
import sys

sys.path.insert(0, str(Path('.tools/python-packages').resolve()))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path('docs/CMS-Fixture-Dataset-Guide.docx')
BLUE = '215C4A'
LIGHT = 'E8F1EE'
GRAY = '5B6470'


def set_font(run, name='Calibri', size=11, bold=False, color='000000'):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:type'), 'dxa')
    tc_w.set(qn('w:w'), str(dxa))


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:type'), 'dxa')
    tbl_ind.set(qn('w:w'), '120')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    shade_paragraph = OxmlElement('w:shd')
    shade_paragraph.set(qn('w:fill'), 'F2F4F7')
    p._p.get_or_add_pPr().append(shade_paragraph)
    run = p.add_run(text)
    set_font(run, 'Consolas', 9.5)


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_font(run)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after, color in [
        ('Heading 1', 16, 18, 10, BLUE),
        ('Heading 2', 13, 14, 7, BLUE),
        ('Heading 3', 12, 10, 5, '1F4D78'),
    ]:
        style = doc.styles[name]
        style.font.name = 'Calibri'
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run('SIDHKOFED CMS | Fixture Dataset'), size=9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run('Fictional test data - not for production use'), size=9, color=GRAY)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run('CMS Fixture Dataset Guide'), size=24, bold=True, color=BLUE)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_font(subtitle.add_run('Reset, seed, verify, export, restore, and test the complete fictional SIDHKOFED dataset.'), size=12, color=GRAY)

    doc.add_heading('Safety And Connection', level=1)
    doc.add_paragraph('All commands are guarded and refuse database names that do not end in _seed. The development database sidhkofed_cms is not modified.')
    table = doc.add_table(rows=3, cols=2)
    set_table_geometry(table, [2700, 6660])
    values = [
        ('Database URL', 'postgresql://postgres:sa3421@127.0.0.1:5432/sidhkofed_seed?schema=public'),
        ('Media root', 'storage/seed-fixtures'),
        ('CMS login', 'Existing SEED_SUPERADMIN_* values from .env'),
    ]
    for i, (label, value) in enumerate(values):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        shade(table.cell(i, 0), LIGHT)
        for run in table.cell(i, 0).paragraphs[0].runs:
            set_font(run, bold=True, color=BLUE)
        for run in table.cell(i, 1).paragraphs[0].runs:
            set_font(run, size=9.5 if i == 0 else 11)

    doc.add_heading('Create, Reset, And Verify', level=1)
    add_code(doc, 'npm run db:fixtures:setup\nnpm run db:fixtures:verify')
    doc.add_paragraph('To discard and rebuild only the isolated fixture database:')
    add_code(doc, 'npm run db:fixtures:reset')
    doc.add_paragraph('The seed is idempotent: stable fixture IDs and upserts prevent duplicate records.')

    doc.add_heading('Run The API With Fixtures', level=1)
    add_code(doc, "$env:DATABASE_URL='postgresql://postgres:sa3421@127.0.0.1:5432/sidhkofed_seed?schema=public'\n$env:STORAGE_PROVIDER='local'\n$env:STORAGE_LOCAL_ROOT=\"$PWD\\storage\\seed-fixtures\"\nnpm run dev")
    doc.add_paragraph('The Admin CMS remains available at http://localhost:3001 and continues to use /api/v1.')

    doc.add_heading('Export And Restore', level=1)
    add_code(doc, 'npm run db:fixtures:export\nnpm run db:fixtures:restore-verify\nnpm run db:fixtures:restore')
    for text in [
        'sidhkofed-seed.sql.gz - portable PostgreSQL schema and data dump',
        'sidhkofed-seed-media.tar.gz - local fixture media objects',
        'manifest.json - migration, checksums, minimums, and exact row counts',
    ]:
        add_bullet(doc, text)

    doc.add_heading('Dataset Coverage', level=1)
    coverage = [
        ('Media', '40 assets; 10 galleries; 30 images; 10 videos'),
        ('Core content', '12 institutions; 10 programmes; 10 toolkits; 18 documents; 24 events; 12 news'),
        ('Public information', '12 communications; 12 tenders; 12 procurement updates'),
        ('Website', '12 pages; 24 menu items; 12 FAQs; 10 digital services'),
        ('Operations', '20 memberships; 15 enquiries; 36 audit logs'),
        ('Analytics', '13 reports; 13 datasets; 52 metrics'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    set_table_geometry(table, [2400, 6960])
    table.rows[0].cells[0].text = 'Area'
    table.rows[0].cells[1].text = 'Fixture records'
    for cell in table.rows[0].cells:
        shade(cell, BLUE)
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=True, color='FFFFFF')
    for area, records in coverage:
        cells = table.add_row().cells
        cells[0].text = area
        cells[1].text = records
        for run in cells[0].paragraphs[0].runs:
            set_font(run, bold=True)
        for run in cells[1].paragraphs[0].runs:
            set_font(run)
    set_table_geometry(table, [2400, 6960])

    doc.add_heading('Known Schema Limits', level=1)
    limits = [
        'SuccessStory has no implemented Prisma model or mounted route.',
        'Homepage content is aggregated from flagged module records; there is no homepage table.',
        'Search vectors are PostgreSQL-generated and are not directly seeded.',
        'Individual beneficiaries, inventory, transactions, hosted video, and report builders are outside scope.',
        'Only the configured Super Admin is created.',
        'MoUs use Institution, Event, and Document relationships.',
    ]
    for item in limits:
        add_bullet(doc, item)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f'Created {OUTPUT}')


if __name__ == '__main__':
    main()
